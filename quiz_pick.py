#!/usr/bin/python3

import docx
import re
#from quiz_selection import quiz_selection
import random

'''
pick 25 randomly from different types of questions.
No.1  ~ No.10  Architecture rules,   10
No.11 ~ No.40  Coding rules,         30
No.41 ~ No.90  C++11 new features,   50
No.91 ~ No.100 C++14 new features,   10

examples:
arc_rule_num = 5
code_rule_num = 7
cpp11_num = 10
cpp14_num = 3
'''
def quiz_selection(arc_rule_num, code_rule_num, cpp11_num, cpp14_num):
	selected = random.sample(range(1,10),arc_rule_num) \
	    + random.sample(range(11,40),code_rule_num) \
	    + random.sample(range(41,90),cpp11_num) \
	    + random.sample(range(91,100),cpp14_num)
	selected.sort()
	return selected

'''
STEP1：Generate quiz file
'''
#read document
file=docx.Document("code_quiz.docx")
#print("paragraphs:"+str(len(file.paragraphs))) 
file_word = docx.Document()

#for para in file.paragraphs:
#    print(para.text)

# dict
dict_num = 0
dict = {dict_num: ''}

for i in range(len(file.paragraphs)):
    t = file.paragraphs[i].text
    s = re.match(r'^\d+\.\s+',t)
    if s:
        dict_num += 1
        dict[dict_num] = t
    else:
        dict[dict_num] += '\n' + t
    #print(dict_num)
    #print(dict[dict_num])

print('Number of quiz questions: ',dict_num)

#print quiz title and attention
file_word.add_paragraph(dict[0])

#randomly generated quiz questions
arc_rule_num = 5
code_rule_num = 7
cpp11_num = 10
cpp14_num = 3
selected = quiz_selection(arc_rule_num,code_rule_num,cpp11_num,cpp14_num)
print('Randomly generated questions:')
print('Architecture rules: {:2} out of 10'.format(arc_rule_num))
print('Coding rules:       {:2} out of 30'.format(code_rule_num))
print('C++11 :             {:2} out of 50'.format(cpp11_num))
print('C++14 :             {:2} out of 10'.format(cpp14_num))
print(selected)

#question_num is the new generated questions number
j=1
for key in selected:
    #print(dict[key])
    #print('\n')
    question_num = str(j) + '. '
    file_word.add_paragraph(question_num + dict[key][dict[key].index('.')+1:])
    j += 1

#create the selected quiz file
print("Generated quiz file: quiz_selected.docx")
file_word.save("quiz_selected.docx")

'''
STEP2：Generate answer file
'''
#read answer document
answer=docx.Document("code_quiz_answer.docx")
answer_word = docx.Document()

# for para in answer.paragraphs:
#    print(para.text)

dict_answer_num = 1
dict_answer = {dict_answer_num: ''}

for i in range(len(answer.paragraphs)):
    t = answer.paragraphs[i].text
    s = re.match(r'^\d+\.\s+',t)
    if s:
        dict_answer[dict_answer_num] = t.strip()
        dict_answer_num += 1

# print(dict_answer)

# selected=[2, 3, 5, 7, 9, 12, 17, 22, 25, 29, 34, 38, 43, 44, 48, 62, 63, 65, 66, 75, 82, 89, 94, 96, 100]

print('\nSelected answer')
j = 1
for key in selected:
    # print(j, dict_answer[key].split('.')[1])
    answer_word.add_paragraph(str(j) + dict_answer[key].split('.')[1])
    j += 1

#create the answer of selected quiz
print("Generated answer file: quiz_selected_answer.docx")
answer_word.save("quiz_selected_answer.docx")